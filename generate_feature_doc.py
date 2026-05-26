from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, color="D0D5DD", size="4", space="0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        edge_el = OxmlElement(f'w:{edge}')
        edge_el.set(qn('w:val'), 'single')
        edge_el.set(qn('w:sz'), size)
        edge_el.set(qn('w:space'), space)
        edge_el.set(qn('w:color'), color)
        tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x24, 0x63)
        run.font.bold = True
    return p

def add_table_custom(doc, headers, rows, header_color="0A2463"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], header_color)
        set_cell_border(hdr_cells[i])
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(10)
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            set_cell_border(row_cells[i])
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if i == 3:
                        try:
                            score = int(str(cell_text).split('/')[0])
                            if score >= 9:
                                run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
                                run.font.bold = True
                            elif score >= 8:
                                run.font.color.rgb = RGBColor(0x0A, 0x24, 0x63)
                            elif score >= 7:
                                run.font.color.rgb = RGBColor(0xC9, 0xA2, 0x27)
                            else:
                                run.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
                        except:
                            pass
    
    doc.add_paragraph()
    return table

doc = Document()

title = doc.add_heading('SecureWealth Twin - Complete Feature Inventory', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x0A, 0x24, 0x63)
    run.font.bold = True
    run.font.size = Pt(22)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('PSB Bank - Govt. of India Undertaking\nHackathon Submission - 2025')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x70, 0x85)
run.font.italic = True

doc.add_paragraph()

add_heading_custom(doc, 'Overall Statistics', level=1)
stats = [
    ('Core Banking & Dashboard', '8', '8.1/10'),
    ('UPI Payments', '9', '8.6/10'),
    ('Rewards & Gamification', '10', '7.9/10'),
    ('Security Beast (10 layers)', '11', '9.1/10'),
    ('Innovation Lab (5 world-firsts)', '5', '9.6/10'),
    ('Protection Suite', '12', '7.9/10'),
    ('Privacy & Compliance', '4', '7.5/10'),
    ('Financial Tools', '6', '7.7/10'),
    ('Gamification & Social', '4', '7.5/10'),
    ('Accessibility & Multi-Mode', '7', '8.4/10'),
    ('AI & Intelligence', '12', '8.3/10'),
    ('PSB Branding & Trust', '6', '8.7/10'),
    ('Utilities & Extras', '11', '7.7/10'),
    ('TOTAL', '105', '8.2/10'),
]
add_table_custom(doc, ['Category', 'Feature Count', 'Average Score'], stats, header_color="1B5E20")

doc.add_page_break()

sections = [
    ('Core Banking & Dashboard', [
        ('Dashboard', 'Net worth, savings, expenses, health score with voice narration', '9/10'),
        ('Wealth Twin', 'AI-powered financial digital twin with predictive modeling', '9/10'),
        ('Goals Tracker', 'SMART goal creation with progress bars and deadlines', '8/10'),
        ('Portfolio Manager', 'Multi-asset portfolio view with allocation charts', '8/10'),
        ('Asset Manager', 'Manual assets + Account Aggregator (AA) linked accounts', '8/10'),
        ('Market View', 'Live market data, NIFTY PE, repo rate, inflation, gold prices', '7/10'),
        ('Scenario Simulator', 'What-if financial scenario modeling', '8/10'),
        ('Add Salary', 'Track salary history and income updates', '7/10'),
    ]),
    ('UPI Payment Simulator', [
        ('Send Money', 'UPI ID / mobile number payment with validation', '9/10'),
        ('Request Money', 'Create payment requests with due dates', '8/10'),
        ('Scan QR', 'Mock QR scanner with auto-fill for merchants', '8/10'),
        ('Pay to Account', 'Account number + IFSC validation and transfer', '8/10'),
        ('Voice Payment', 'Web Speech API - Pay 500 to Rahul', '9/10'),
        ('MPIN Entry', '6-digit numeric keypad with masking', '9/10'),
        ('Transaction History', 'Persistent localStorage with CSV export', '8/10'),
        ('Quick Pay Card', 'Dashboard-integrated instant payment widget', '9/10'),
        ('Payment Hub Bar', 'Sticky persistent payment bar on every page', '9/10'),
    ]),
    ('Rewards & Gamification', [
        ('Cashback Engine', 'Merchant-funded cashback (Swiggy, Amazon, Zomato)', '9/10'),
        ('Cashback Piggy Bank', 'Visual SVG piggy that fills with earnings', '9/10'),
        ('Login Streak', 'Daily streak with 7/30/100 day milestones', '8/10'),
        ('Spin Wheel', 'Post-payment mini-game for extra 1-5% cashback', '9/10'),
        ('Mood Meter', 'Emotion tracking with good-habit rewards', '8/10'),
        ('Bill Splitter', 'AI-powered equal/percentage/custom split', '8/10'),
        ('Payment Requests', 'Social payment requests with pay button', '7/10'),
        ('Group Savings Jars', 'Shared goals with contribution tracking', '7/10'),
        ('Referral Rewards', '50 each for friend signup (CAC-funded)', '7/10'),
        ('Ad Rewards', 'Watch simulated ad for 2 cashback', '7/10'),
    ]),
    ('Security Beast - 10 Layers', [
        ('TPM Attestation', 'Hardware root-of-trust simulation', '9/10'),
        ('eBPF Monitor', 'Kernel-level real-time security monitoring', '9/10'),
        ('Honeytoken Manager', 'Fake credentials to detect intruders', '9/10'),
        ('FIDO2 Passkey Auth', 'Passwordless biometric authentication', '9/10'),
        ('Post-Quantum Crypto', 'CRYSTALS-Kyber lattice-based encryption', '10/10'),
        ('Behavioral Biometrics', 'Typing pattern and behavior analysis', '9/10'),
        ('Decentralized ID (DID)', 'Self-sovereign identity with QR codes', '9/10'),
        ('Transaction Trap', 'Fake OTP phishing honeypot for scammers', '9/10'),
        ('Secure Enclave', 'Hardware security module attestation', '9/10'),
        ('Blockchain Audit', 'SHA-256 tamper-evident audit chain', '9/10'),
        ('Trust Score (0-100)', 'Aggregated live security score badge', '9/10'),
    ]),
    ('Innovation Lab - 5 World-Firsts', [
        ('Parametric Insurance', 'Auto-payouts for flight delay, border, weather', '10/10'),
        ('Ghost Mode', 'AI bot that traps and wastes scammer time', '10/10'),
        ('Dead Mans Switch', 'Auto-inheritance if check-in missed', '10/10'),
        ('Gig Income Smoother', 'Synthetic stable salary from irregular income', '9/10'),
        ('Social Collateral Loans', 'Friend-vouch based credit system', '9/10'),
    ]),
    ('Protection Suite', [
        ('Risk Meter', 'Real-time fraud risk scoring', '8/10'),
        ('Fraud Simulator', 'Interactive attack simulation', '8/10'),
        ('Panic Button', 'Instant account freeze + alert', '8/10'),
        ('Duress Mode', 'Fake sanitized view under coercion', '8/10'),
        ('Duress PIN', 'Separate PIN that triggers silent alert', '8/10'),
        ('Threat Intel', 'Latest fraud patterns and IOCs', '7/10'),
        ('OTP Simulation', 'Secure OTP flow demo', '7/10'),
        ('Secure Checkout', 'Multi-factor payment verification', '8/10'),
        ('Payment Guard', 'Real-time fraud database check', '8/10'),
        ('Cooling-Off Vault', '30-second delay for high-value transactions', '8/10'),
        ('Scam Caller ID', 'Identify and block scam calls', '8/10'),
        ('Stress Test', 'Portfolio resilience under market shocks', '8/10'),
    ]),
    ('Privacy & Compliance', [
        ('Privacy Center', 'Data consent management and scope control', '8/10'),
        ('Consent Manager', 'Granular data sharing permissions', '8/10'),
        ('Compliance Badges', 'RBI, SEBI, IRDAI, NPCI certification display', '7/10'),
        ('Demo Controls', 'Judge mode quick-access toggles', '7/10'),
    ]),
    ('Financial Tools', [
        ('Tax Planner', 'Tax optimization recommendations', '7/10'),
        ('Financial Calculators', 'SIP, EMI, FD, RD, PPF, NPS calculators', '8/10'),
        ('Bill Calendar', 'Recurring bill tracking with predictions', '8/10'),
        ('Credit Health (CIBIL)', 'Credit score simulation with factor breakdown', '8/10'),
        ('Digital Gold', 'Gold investment tracker', '7/10'),
        ('Account Aggregator', 'RBI AA framework account linking demo', '8/10'),
    ]),
    ('Gamification & Social', [
        ('Fantasy League', 'Stock-picking competition with leaderboard', '8/10'),
        ('Boosts Manager', 'Financial habit boosters and power-ups', '7/10'),
        ('Investment Quiz', 'Risk profiling and knowledge assessment', '8/10'),
        ('Badge Streak', 'Achievement badges for milestones', '7/10'),
    ]),
    ('Accessibility & Multi-Mode', [
        ('Dark Mode', 'Full dark theme support', '9/10'),
        ('Senior Mode', 'Large fonts, simplified UI, high contrast', '9/10'),
        ('Kids Mode', 'Smart piggy bank with tasks and allowance', '9/10'),
        ('NRI Mode', 'Dual currency (INR/USD), NRE/NRO accounts', '8/10'),
        ('Business Mode', 'GST, invoicing, business analytics', '8/10'),
        ('Family View', 'Household wealth aggregation', '8/10'),
        ('Accessibility Settings', 'WCAG 2.1 AA compliant controls', '8/10'),
    ]),
    ('AI & Intelligence', [
        ('NBA Co-Pilot', 'Next Best Action AI recommendations', '9/10'),
        ('AI Decision Log', 'Transparent audit of all AI recommendations', '9/10'),
        ('Financial Twin Chat', 'Conversational AI wealth advisor', '9/10'),
        ('ELI5 Tooltips', 'Explain Like I am 5 for financial terms', '8/10'),
        ('Financial Literacy Cards', 'Bite-sized learning modules', '8/10'),
        ('Wealth DNA', 'Personalized wealth personality analysis', '8/10'),
        ('Wealth Benchmark', 'Peer comparison and percentile ranking', '8/10'),
        ('Financial Weather', 'Market conditions as weather forecast', '8/10'),
        ('Adaptive Insights', 'Context-aware smart suggestions', '8/10'),
        ('Monthly Narrative', 'AI-generated monthly financial story', '8/10'),
        ('Behavioral Nudges', 'Psychology-based spending prompts', '8/10'),
        ('Emotion Check-in', 'Mood-linked financial wellness tracking', '7/10'),
    ]),
    ('PSB Branding & Trust', [
        ('PSB Bank Header', 'Ashoka Chakra logo + Govt. of India branding', '9/10'),
        ('Trust Banner', 'Rotating DICGC/RBI/security messages', '9/10'),
        ('Government Schemes', 'PMJDY, SSY, APY, PMVVY, PMJJBY info', '9/10'),
        ('Accessible Footer', 'RBI Ombudsman, Grievance, Regulatory info', '9/10'),
        ('Security Health Widget', 'Live trust score + active layers display', '8/10'),
        ('Welcome Banner', 'Personalized greeting + DICGC insured badge', '8/10'),
    ]),
    ('Utilities & Extras', [
        ('Demo Credit Card', 'Virtual card with security features demo', '8/10'),
        ('Transaction Comparison', 'AI duplicate detection and merge', '8/10'),
        ('System Architecture', 'Full tech stack visualization', '8/10'),
        ('Report Generator', 'PDF/HTML financial report export', '8/10'),
        ('Pitch Mode', 'Presentation-friendly spotlight view', '7/10'),
        ('Notification Demo', 'Push notification simulation', '7/10'),
        ('Coerced Mode', 'Sanitized view for duress situations', '8/10'),
        ('Lockdown Overlay', 'Full-screen emergency lock', '8/10'),
        ('Biometric Auth', 'Face/touch simulation', '8/10'),
        ('8 Language Support', 'EN, HI, TA, TE, ML, KN, MR, BN', '7/10'),
        ('Voice Narration', 'Screen reader with Indian number words', '8/10'),
    ]),
]

for section_title, features in sections:
    add_heading_custom(doc, section_title, level=1)
    rows = [(str(i+1), name, desc, score) for i, (name, desc, score) in enumerate(features)]
    add_table_custom(doc, ['#', 'Feature', 'Description', 'Score'], rows)

doc.save('C:/Users/sdeep/Final PSB/SecureWealth_Twin_Feature_Scorecard.docx')
print('Document saved successfully!')
